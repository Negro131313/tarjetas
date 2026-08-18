import os
import re
import traceback
from docx import Document

# ============================================================
# ARCHIVO (RUTA CORREGIDA)
# ============================================================
ARCHIVO = r"C:\Users\Julio\Desktop\rembg_project\literatura\cambiar.docx"

if not os.path.exists(ARCHIVO):
    print("============================================")
    print(" ERROR: EL ARCHIVO NO EXISTE")
    print("============================================")
    print(f"Ruta no encontrada:\n{ARCHIVO}")
else:
    try:
        print("============================================")
        print(" ARCHIVO ENCONTRADO - PROCESANDO FORMATO")
        print("============================================")

        doc = Document(ARCHIVO)

        # 1. Extraer todas las líneas del documento original
        lineas = [p.text for p in doc.paragraphs]
        texto_completo = "\n".join(lineas)

        # 2. Limpiar tabulaciones y caracteres invisibles
        texto_completo = texto_completo.replace("\xa0", " ").replace("\t", " ")

        # 3. Dividir por saltos dobles o múltiples de línea (bloques principales)
        bloques = re.split(r"\n\s*\n+", texto_completo)

        bloques_procesados = []
        for bloque in bloques:
            # Eliminar todos los saltos de línea dentro del bloque y cambiarlos por un espacio
            texto_unido = re.sub(r"\n+", " ", bloque)

            # Insertar un espacio si una cita terminada en corchete o punto quedó pegada a la siguiente palabra
            # Ej: "Jefferson.[76]Identifying" -> "Jefferson.[76] Identifying"
            texto_unido = re.sub(r"(\]|\.)([A-ZÁÉÍÓÚÑa-z])", r"\1 \2", texto_unido)

            # Reducir espacios dobles a un solo espacio
            texto_limpio = re.sub(r"\s+", " ", texto_unido).strip()

            if texto_limpio:
                bloques_procesados.append(texto_limpio)

        # 4. Volver a escribir el documento con los bloques exactos
        doc_nuevo = Document()
        for b in bloques_procesados:
            doc_nuevo.add_paragraph(b)

        doc_nuevo.save(ARCHIVO)

        print("============================================")
        print(" ¡PROCESO COMPLETADO EXITOSAMENTE!")
        print("============================================")
        print("Se unieron los saltos internos y se separaron los bloques pegados.")

    except PermissionError:
        print("\n[ERROR DE PERMISOS] Cierra Microsoft Word antes de ejecutar el código.")
    except Exception as e:
        print("\n============================================")
        print(" ERROR DETECTADO")
        print("============================================")
        print(e)
        traceback.print_exc()

input("\nPresiona ENTER para cerrar...")